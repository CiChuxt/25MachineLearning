"""
生成机器学习课程论文Word文档（格式修订版）
格式要求：
- 一级标题：三号黑体不加粗
- 二级标题：四号黑体不加粗
- 三级标题：小四宋体不加粗
- 正文：小四宋体不加粗
- 段前段后为0
- 行距1.5倍
- 正文首行缩进两个字符
- 两端对齐
- 公式使用LaTeX渲染
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import os

# 路径设置
BASE_DIR = r'C:\Users\航仔的Computer\Desktop\机器学习论文_代码数据'
FIG_DIR = os.path.join(BASE_DIR, 'figures')
FORMULA_DIR = os.path.join(BASE_DIR, 'formulas')
OUTPUT_PATH = r'C:\Users\航仔的Computer\Desktop\基于PCA与t-SNE降维的手写数字识别性能对比研究.docx'

os.makedirs(FORMULA_DIR, exist_ok=True)


def render_latex_formula(latex_str, filename, fontsize=14):
    """使用matplotlib渲染LaTeX公式为图片"""
    fig = plt.figure(figsize=(6, 1))
    fig.text(0.5, 0.5, f'${latex_str}$', fontsize=fontsize, 
             ha='center', va='center')
    plt.axis('off')
    save_path = os.path.join(FORMULA_DIR, filename)
    plt.savefig(save_path, dpi=200, bbox_inches='tight', pad_inches=0.1,
                transparent=True)
    plt.close()
    return save_path


def set_paragraph_format(p, first_line_indent=True, alignment='justify'):
    """设置段落格式：段前段后0，行距1.5倍，首行缩进，两端对齐"""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 小四12pt，2字符=24pt
    
    if alignment == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_font(run, font_name='宋体', size=12, bold=False, english_font='Times New Roman'):
    """设置字体，中文宋体，英文Times New Roman
    小四=12pt, 四号=14pt, 三号=16pt
    """
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = english_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level=1):
    """添加标题
    level=1: 三号黑体不加粗 (16pt)
    level=2: 四号黑体不加粗 (14pt)
    level=3: 小四宋体不加粗 (12pt)
    """
    heading = doc.add_paragraph()
    set_paragraph_format(heading, first_line_indent=False, alignment='left')
    
    run = heading.add_run(text)
    if level == 1:
        set_font(run, font_name='黑体', size=16, bold=False)
    elif level == 2:
        set_font(run, font_name='黑体', size=14, bold=False)
    elif level == 3:
        set_font(run, font_name='宋体', size=12, bold=False)
    
    return heading


def add_paragraph(doc, text, indent=True):
    """添加正文段落：小四宋体不加粗，首行缩进，两端对齐，1.5倍行距"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=indent, alignment='justify')
    run = p.add_run(text)
    set_font(run, font_name='宋体', size=12, bold=False)
    return p


def add_formula(doc, latex_str, formula_name):
    """添加LaTeX公式（居中）"""
    img_path = render_latex_formula(latex_str, f'{formula_name}.png')
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=False, alignment='center')
    run = p.add_run()
    run.add_picture(img_path, height=Inches(0.35))
    return p


def add_figure(doc, img_path, caption):
    """添加图片和图题（图题在图下方）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=False, alignment='center')
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    
    # 图题
    caption_p = doc.add_paragraph()
    set_paragraph_format(caption_p, first_line_indent=False, alignment='center')
    run = caption_p.add_run(caption)
    set_font(run, font_name='宋体', size=12, bold=False)


def add_table(doc, data, headers, caption):
    """添加表格和表标题（表标题在表上方）"""
    # 表标题
    caption_p = doc.add_paragraph()
    set_paragraph_format(caption_p, first_line_indent=False, alignment='center')
    run = caption_p.add_run(caption)
    set_font(run, font_name='宋体', size=12, bold=False)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, font_name='宋体', size=12, bold=True)
    
    # 数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_data))
            set_font(run, font_name='宋体', size=12, bold=False)
    
    doc.add_paragraph()  # 空行


def generate_all_formulas():
    """预渲染所有公式"""
    formulas = {
        'pca_objective': r'\max_{\mathbf{w}} \mathbf{w}^T \Sigma \mathbf{w}, \quad s.t. \quad \mathbf{w}^T \mathbf{w} = 1',
        'pca_lagrange': r'L(\mathbf{w}, \lambda) = \mathbf{w}^T \Sigma \mathbf{w} - \lambda(\mathbf{w}^T \mathbf{w} - 1)',
        'pca_eigen': r'\Sigma \mathbf{w} = \lambda \mathbf{w}',
        'pca_cumulative': r'\text{Cumulative Variance Ratio} = \frac{\sum_{i=1}^{k} \lambda_i}{\sum_{i=1}^{d} \lambda_i}',
        'tsne_pji': r'p_{j|i} = \frac{\exp(-\|\mathbf{x}_i - \mathbf{x}_j\|^2 / 2\sigma_i^2)}{\sum_{k \neq i} \exp(-\|\mathbf{x}_i - \mathbf{x}_k\|^2 / 2\sigma_i^2)}',
        'tsne_perp': r'\text{Perp}(P_i) = 2^{H(P_i)}, \quad H(P_i) = -\sum_j p_{j|i} \log_2 p_{j|i}',
        'tsne_pij': r'p_{ij} = \frac{p_{j|i} + p_{i|j}}{2n}',
        'tsne_qij': r'q_{ij} = \frac{(1 + \|\mathbf{y}_i - \mathbf{y}_j\|^2)^{-1}}{\sum_{k \neq l} (1 + \|\mathbf{y}_k - \mathbf{y}_l\|^2)^{-1}}',
        'tsne_kl': r'C = KL(P\|Q) = \sum_i \sum_j p_{ij} \log \frac{p_{ij}}{q_{ij}}',
        'tsne_gradient': r'\frac{\partial C}{\partial \mathbf{y}_i} = 4 \sum_j (p_{ij} - q_{ij}) (\mathbf{y}_i - \mathbf{y}_j) (1 + \|\mathbf{y}_i - \mathbf{y}_j\|^2)^{-1}',
        'zscore': r"x' = \frac{x - \mu}{\sigma}",
    }
    
    for name, latex in formulas.items():
        render_latex_formula(latex, f'{name}.png', fontsize=14)
    print('所有公式渲染完成')


def generate_paper():
    """生成完整论文"""
    # 先生成所有公式
    generate_all_formulas()
    
    doc = Document()
    
    # 设置默认样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.size = Pt(12)
    
    # ========== 封面 ==========
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False)
    
    title_p = doc.add_paragraph()
    set_paragraph_format(title_p, first_line_indent=False, alignment='center')
    run = title_p.add_run('基于PCA与t-SNE降维的手写数字识别性能对比研究')
    set_font(run, font_name='黑体', size=22, bold=True)  # 封面标题大一些
    
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False)
    
    # 课程信息
    info_lines = [
        ('课程名称：', '机器学习'),
        ('学生姓名：', '刘宇航'),
        ('学    号：', '202521511102'),
        ('学    院：', '数学与计算科学学院'),
        ('专    业：', '应用统计'),
    ]
    
    for label, value in info_lines:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False, alignment='center')
        run1 = p.add_run(label)
        set_font(run1, font_name='宋体', size=14, bold=False)
        run2 = p.add_run(value)
        set_font(run2, font_name='宋体', size=14, bold=False)
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_heading(doc, '目录', level=1)
    
    toc_items = [
        '1 研究背景',
        '  1.1 高维数据与维数灾难',
        '  1.2 降维技术研究现状',
        '  1.3 研究意义与主要内容',
        '2 研究方法',
        '  2.1 主成分分析（PCA）',
        '    2.1.1 方法提出动机',
        '    2.1.2 目标函数与数学推导',
        '    2.1.3 算法流程',
        '  2.2 t分布随机邻域嵌入（t-SNE）',
        '    2.2.1 方法提出动机',
        '    2.2.2 目标函数与数学推导',
        '    2.2.3 算法流程',
        '  2.3 分类器选择',
        '3 数据描述',
        '  3.1 数据集来源',
        '  3.2 数据特点分析',
        '  3.3 数据预处理',
        '4 结果对比与分析',
        '  4.1 实验设置',
        '  4.2 降维可视化结果分析',
        '  4.3 分类准确率对比分析',
        '  4.4 计算效率对比分析',
        '5 总结',
        '参考文献',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False, alignment='left')
        run = p.add_run(item)
        set_font(run, font_name='宋体', size=12, bold=False)
    
    doc.add_page_break()
    
    # ========== 1 研究背景 ==========
    add_heading(doc, '1 研究背景', level=1)
    
    add_heading(doc, '1.1 高维数据与维数灾难', level=2)
    add_paragraph(doc, '随着信息技术的快速发展，各领域产生的数据规模呈现爆炸式增长，数据的维度也随之急剧增加。在图像识别、文本分类、生物信息学等领域，数据维度往往达到成百上千维，甚至更高。高维数据在带来丰富信息的同时，也给传统的机器学习算法带来了严峻挑战，这一现象被称为"维数灾难"（Curse of Dimensionality）。')
    add_paragraph(doc, '维数灾难主要表现在以下几个方面：首先，随着维度增加，数据样本在高维空间中变得极其稀疏，导致基于距离的分类算法（如K近邻）性能急剧下降；其次，高维空间中所有样本点之间的距离趋于相等，使得距离度量失去了区分度；再次，高维特征中往往存在大量冗余信息和噪声，增加了模型的计算复杂度和过拟合风险；最后，高维数据的可视化极其困难，人类只能直观理解三维及以下的数据空间。')
    
    add_heading(doc, '1.2 降维技术研究现状', level=2)
    add_paragraph(doc, '为了解决维数灾难问题，降维（Dimensionality Reduction）技术应运而生。降维旨在将高维数据映射到低维空间，同时尽可能保留原始数据的重要结构信息。根据映射方式的不同，降维方法大致可以分为线性降维和非线性降维两大类。')
    add_paragraph(doc, '线性降维方法中，主成分分析（Principal Component Analysis, PCA）是最经典、应用最广泛的方法之一，由Pearson于1901年提出，后由Hotelling于1933年进一步发展。PCA通过正交变换将原始高维特征转换为一组线性不相关的主成分，在最大化方差的前提下实现数据降维。此外，线性判别分析（Linear Discriminant Analysis, LDA）、多维缩放（Multidimensional Scaling, MDS）等也是经典的线性降维方法。')
    add_paragraph(doc, '然而，现实世界中的许多数据本质上具有非线性结构，线性降维方法难以捕捉数据的内在流形结构。为此，研究者们提出了一系列非线性降维方法，即流形学习方法。代表性方法包括：Tenenbaum等人于2000年提出的等距映射（Isomap），Roweis和Saul于2000年提出的局部线性嵌入（Locally Linear Embedding, LLE），以及Hinton和Roweis于2002年提出的随机邻域嵌入（Stochastic Neighbor Embedding, SNE）。')
    add_paragraph(doc, '在SNE的基础上，van der Maaten和Hinton于2008年提出了t分布随机邻域嵌入（t-distributed Stochastic Neighbor Embedding, t-SNE）。t-SNE通过在低维空间使用t分布替代高斯分布，有效解决了SNE中的拥挤问题，成为目前最流行的高维数据可视化工具之一，在机器学习、数据挖掘、生物信息学等领域得到了广泛应用。')
    
    add_heading(doc, '1.3 研究意义与主要内容', level=2)
    add_paragraph(doc, '尽管降维技术已经取得了长足发展，但不同降维方法在不同数据集和任务上的表现差异显著。线性方法计算高效但难以处理非线性结构，非线性方法可视化效果好但计算复杂度高。系统对比不同降维方法的性能特点，对于实际应用中的算法选择具有重要的指导意义。')
    add_paragraph(doc, '本文选取两种具有代表性的降维方法——线性降维的代表PCA和非线性降维的代表t-SNE，以手写数字识别任务为例，从分类准确率、可视化效果、计算效率三个维度进行系统的对比实验研究。本文的主要内容包括：首先介绍PCA和t-SNE的理论基础与算法原理；其次介绍实验所用的数据集和预处理方法；然后通过对比实验分析两种降维方法的优缺点；最后总结全文并展望未来研究方向。')
    
    doc.add_page_break()
    
    # ========== 2 研究方法 ==========
    add_heading(doc, '2 研究方法', level=1)
    
    add_heading(doc, '2.1 主成分分析（PCA）', level=2)
    
    add_heading(doc, '2.1.1 方法提出动机', level=3)
    add_paragraph(doc, '在高维数据中，不同特征之间往往存在相关性，导致信息冗余。PCA的核心思想是通过线性变换，将原始的高维相关特征转换为少数几个互不相关的综合指标（主成分），这些主成分能够尽可能多地保留原始数据的变异信息。其提出动机主要包括：（1）消除特征间的相关性，减少信息冗余；（2）在保留主要信息的前提下降低数据维度，提高后续处理效率；（3）通过提取主成分来发现数据的主要变化方向，揭示数据的内在结构。')
    
    add_heading(doc, '2.1.2 目标函数与数学推导', level=3)
    add_paragraph(doc, '设原始数据矩阵为 X ∈ R^(n×d)，其中n为样本数，d为特征维度。假设数据已经过中心化处理，即每个特征的均值为0。PCA的目标是寻找一组正交基向量（主成分方向），使得数据在这些方向上的投影方差最大。')
    
    add_paragraph(doc, '第一主成分的优化目标为：', indent=False)
    add_formula(doc, r'\max_{\mathbf{w}} \mathbf{w}^T \Sigma \mathbf{w}, \quad s.t. \quad \mathbf{w}^T \mathbf{w} = 1', 'pca_objective')
    
    add_paragraph(doc, '其中 Σ 为数据的协方差矩阵，w 为投影方向。使用拉格朗日乘数法求解：')
    add_formula(doc, r'L(\mathbf{w}, \lambda) = \mathbf{w}^T \Sigma \mathbf{w} - \lambda(\mathbf{w}^T \mathbf{w} - 1)', 'pca_lagrange')
    
    add_paragraph(doc, '对w求导并令其为0，可得：')
    add_formula(doc, r'\Sigma \mathbf{w} = \lambda \mathbf{w}', 'pca_eigen')
    
    add_paragraph(doc, '这表明w是协方差矩阵Σ的特征向量，对应的特征值λ即为投影后方差。因此，第一主成分方向就是协方差矩阵最大特征值对应的特征向量。第k主成分方向则是第k大特征值对应的特征向量，且与前k-1个主成分方向正交。')
    
    add_paragraph(doc, '设前k个主成分对应的特征值为 λ₁ ≥ λ₂ ≥ ... ≥ λ_k，则前k个主成分的累计方差解释率为：')
    add_formula(doc, r'\text{Cumulative Variance Ratio} = \frac{\sum_{i=1}^{k} \lambda_i}{\sum_{i=1}^{d} \lambda_i}', 'pca_cumulative')
    
    add_heading(doc, '2.1.3 算法流程', level=3)
    add_paragraph(doc, 'PCA算法的具体流程如下：', indent=False)
    steps = [
        '步骤1：对原始数据进行中心化处理，即每个特征减去其均值；',
        '步骤2：计算数据的协方差矩阵 Σ = (1/n) X^T X；',
        '步骤3：对协方差矩阵进行特征值分解，得到特征值和对应的特征向量；',
        '步骤4：将特征值按从大到小排序，选取前k个最大的特征值对应的特征向量；',
        '步骤5：将这k个特征向量组成投影矩阵 W ∈ R^(d×k)；',
        '步骤6：将原始数据投影到低维空间：Z = XW，得到降维后的数据 Z ∈ R^(n×k)。'
    ]
    for step in steps:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False, alignment='justify')
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(step)
        set_font(run, font_name='宋体', size=12, bold=False)
    
    add_heading(doc, '2.2 t分布随机邻域嵌入（t-SNE）', level=2)
    
    add_heading(doc, '2.2.1 方法提出动机', level=3)
    add_paragraph(doc, '传统的线性降维方法如PCA主要关注数据的全局结构，对于具有非线性流形结构的数据往往效果不佳。t-SNE作为一种非线性降维方法，其核心动机是在低维空间中保持高维空间中数据点之间的局部邻域关系。与SNE相比，t-SNE的提出主要解决了两个问题：一是SNE的对称化版本优化困难，二是低维空间中的"拥挤问题"（Crowding Problem）——即高维空间中中等距离的点在低维空间中被迫挤在一起。')
    add_paragraph(doc, 't-SNE的创新之处在于：在高维空间使用高斯分布计算点对之间的相似度，而在低维空间使用自由度为1的t分布（即柯西分布）来建模相似度。由于t分布具有 heavier tail 的特性，能够更好地在低维空间中拉开不同簇之间的距离，从而有效缓解拥挤问题，产生更优的可视化效果。')
    
    add_heading(doc, '2.2.2 目标函数与数学推导', level=3)
    add_paragraph(doc, 't-SNE的基本思想是将高维空间中点对之间的相似性转化为条件概率，然后在低维空间中寻找一种表示，使得对应的概率分布尽可能相似。')
    
    add_paragraph(doc, '首先，在高维空间中，数据点x_i与x_j之间的相似度用条件概率p_{j|i}表示：', indent=False)
    add_formula(doc, r'p_{j|i} = \frac{\exp(-\|\mathbf{x}_i - \mathbf{x}_j\|^2 / 2\sigma_i^2)}{\sum_{k \neq i} \exp(-\|\mathbf{x}_i - \mathbf{x}_k\|^2 / 2\sigma_i^2)}', 'tsne_pji')
    
    add_paragraph(doc, '其中σ_i是以x_i为中心的高斯分布的方差，通过困惑度（Perplexity）参数来调节。困惑度定义为：')
    add_formula(doc, r'\text{Perp}(P_i) = 2^{H(P_i)}, \quad H(P_i) = -\sum_j p_{j|i} \log_2 p_{j|i}', 'tsne_perp')
    
    add_paragraph(doc, '为了对称化处理，定义联合概率分布：')
    add_formula(doc, r'p_{ij} = \frac{p_{j|i} + p_{i|j}}{2n}', 'tsne_pij')
    
    add_paragraph(doc, '在低维空间中，对应的点为y_i和y_j，使用自由度为1的t分布计算相似度q_{ij}：')
    add_formula(doc, r'q_{ij} = \frac{(1 + \|\mathbf{y}_i - \mathbf{y}_j\|^2)^{-1}}{\sum_{k \neq l} (1 + \|\mathbf{y}_k - \mathbf{y}_l\|^2)^{-1}}', 'tsne_qij')
    
    add_paragraph(doc, 't-SNE的目标是最小化高维分布P和低维分布Q之间的KL散度：')
    add_formula(doc, r'C = KL(P\|Q) = \sum_i \sum_j p_{ij} \log \frac{p_{ij}}{q_{ij}}', 'tsne_kl')
    
    add_paragraph(doc, '使用梯度下降法优化该目标函数。对y_i求梯度可得：')
    add_formula(doc, r'\frac{\partial C}{\partial \mathbf{y}_i} = 4 \sum_j (p_{ij} - q_{ij}) (\mathbf{y}_i - \mathbf{y}_j) (1 + \|\mathbf{y}_i - \mathbf{y}_j\|^2)^{-1}', 'tsne_gradient')
    
    add_heading(doc, '2.2.3 算法流程', level=3)
    add_paragraph(doc, 't-SNE算法的具体流程如下：', indent=False)
    steps = [
        '步骤1：设定困惑度Perp、迭代次数、学习率等超参数；',
        '步骤2：对每个数据点x_i，通过二分搜索确定σ_i，使得高斯分布的困惑度等于设定值；',
        '步骤3：计算高维空间中的联合概率分布p_{ij}；',
        '步骤4：随机初始化低维表示Y（通常从正态分布中采样）；',
        '步骤5：计算低维空间中的联合概率分布q_{ij}；',
        '步骤6：计算KL散度损失及其梯度；',
        '步骤7：使用梯度下降（带动量）更新低维表示Y；',
        '步骤8：重复步骤5-7，直到达到最大迭代次数或收敛。'
    ]
    for step in steps:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False, alignment='justify')
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(step)
        set_font(run, font_name='宋体', size=12, bold=False)
    
    add_heading(doc, '2.3 分类器选择', level=2)
    add_paragraph(doc, '为了全面评估降维方法对分类任务的影响，本文选取三种经典的分类器进行对比实验，分别是支持向量机（SVM）、K近邻（KNN）和随机森林（Random Forest）。')
    add_paragraph(doc, '支持向量机通过寻找最大间隔超平面实现分类，使用核技巧可以处理非线性问题，具有良好的泛化能力。本文选用RBF核的SVM作为基准分类器之一。K近邻是一种基于实例的学习算法，通过投票方式进行分类，其性能高度依赖于特征空间的距离度量，因此对降维效果较为敏感。随机森林是一种集成学习方法，通过构建多棵决策树并投票得到最终结果，具有较强的抗过拟合能力。')
    
    doc.add_page_break()
    
    # ========== 3 数据描述 ==========
    add_heading(doc, '3 数据描述', level=1)
    
    add_heading(doc, '3.1 数据集来源', level=2)
    add_paragraph(doc, '本文实验使用scikit-learn机器学习库内置的手写数字数据集（Digits Dataset）。该数据集来源于UCI机器学习仓库，由Alpaydin和Kaynak于1998年整理发布。原始数据由43人书写的数字扫描而成，经过归一化处理后得到8×8的灰度图像。')
    add_paragraph(doc, '数据集官方链接：https://scikit-learn.org/stable/datasets/toy_dataset.html#digits-dataset')
    add_paragraph(doc, '原始UCI数据集链接：https://archive.ics.uci.edu/ml/datasets/Pen-Based+Recognition+of+Handwritten+Digits')
    
    add_heading(doc, '3.2 数据特点分析', level=2)
    add_paragraph(doc, '该数据集具有以下特点：')
    
    # 数据集基本信息表
    table_data = [
        ['样本总数', '1797'],
        ['特征维度', '64维（8×8像素）'],
        ['类别数量', '10类（数字0-9）'],
        ['像素取值范围', '0-16（整数灰度值）'],
        ['训练集样本数', '1257（70%）'],
        ['测试集样本数', '540（30%）'],
    ]
    add_table(doc, table_data, ['属性', '数值'], '表3-1 手写数字数据集基本信息')
    
    add_paragraph(doc, '从类别分布来看，各类别样本数量较为均衡，数字0有178个样本，数字1有182个样本，数字8样本最少为174个，不存在严重的类别不平衡问题。每幅图像为8×8的灰度图，像素值范围为0到16的整数，其中0代表白色背景，16代表最黑的笔迹。')
    add_paragraph(doc, '从维度角度看，64维虽然远低于真实场景中的图像数据（如MNIST的784维），但已足以体现降维算法的特性。同时，该数据集规模适中，计算效率高，适合作为教学和算法对比研究的基准数据集。')
    
    add_heading(doc, '3.3 数据预处理', level=2)
    add_paragraph(doc, '为了保证实验的公平性和算法的稳定性，本文对数据进行了以下预处理操作：')
    add_paragraph(doc, '（1）标准化处理：由于PCA和距离相关算法对特征尺度敏感，本文使用Z-score标准化将每个特征缩放到均值为0、标准差为1的标准正态分布。标准化公式为：', indent=False)
    add_formula(doc, r"x' = \frac{x - \mu}{\sigma}", 'zscore')
    add_paragraph(doc, '其中μ为特征均值，σ为特征标准差。')
    add_paragraph(doc, '（2）数据集划分：采用分层抽样（Stratified Sampling）的方式，按照7:3的比例将数据集划分为训练集和测试集，保证训练集和测试集中各类别的比例与原始数据一致。')
    add_paragraph(doc, '（3）随机种子设置：为保证实验可复现，所有随机操作均设置随机种子为42。')
    
    doc.add_page_break()
    
    # ========== 4 结果对比与分析 ==========
    add_heading(doc, '4 结果对比与分析', level=1)
    
    add_heading(doc, '4.1 实验设置', level=2)
    add_paragraph(doc, '本文的实验环境为Python 3.13，使用scikit-learn 1.6机器学习库。实验硬件为普通个人计算机。所有降维方法均将数据降至2维，以便于可视化和公平对比。')
    add_paragraph(doc, '具体参数设置如下：')
    add_paragraph(doc, '（1）PCA：保留2个主成分，使用奇异值分解（SVD）求解；', indent=False)
    add_paragraph(doc, '（2）t-SNE：降维到2维，困惑度perplexity=30，使用PCA初始化，学习率自动调整；', indent=False)
    add_paragraph(doc, '（3）SVM分类器：使用RBF核函数，正则化参数C=1.0；', indent=False)
    add_paragraph(doc, '（4）KNN分类器：近邻数k=5，使用欧氏距离；', indent=False)
    add_paragraph(doc, '（5）随机森林：决策树数量n_estimators=100。', indent=False)
    add_paragraph(doc, '评价指标包括：测试集分类准确率、5折交叉验证准确率、算法运行时间。')
    
    add_heading(doc, '4.2 降维可视化结果分析', level=2)
    add_paragraph(doc, '首先分析两种降维方法的方差解释特性。图4-1展示了PCA各主成分的累计方差解释率曲线。')
    
    # PCA方差解释率图
    if os.path.exists(os.path.join(FIG_DIR, 'pca_variance_explained.png')):
        add_figure(doc, os.path.join(FIG_DIR, 'pca_variance_explained.png'), 
                  '图4-1 PCA累计方差解释率曲线')
    
    add_paragraph(doc, '从图4-1可以看出，前2个主成分仅能解释约22.37%的数据方差。要达到90%的累计方差解释率，大约需要前30个左右的主成分。这说明手写数字数据的信息分布较为分散，仅用前两个主成分难以保留大部分信息，这也是PCA降维到2维后分类性能下降较多的重要原因。')
    
    add_paragraph(doc, '接下来对比两种方法的二维可视化效果。图4-2和图4-3分别展示了PCA和t-SNE将训练集数据降至2维后的散点图，不同颜色代表不同的数字类别。')
    
    # PCA可视化图
    if os.path.exists(os.path.join(FIG_DIR, 'pca_2d_visualization.png')):
        add_figure(doc, os.path.join(FIG_DIR, 'pca_2d_visualization.png'),
                  '图4-2 PCA二维降维可视化结果（训练集）')
    
    add_paragraph(doc, '从图4-2可以看出，PCA降维后不同数字的样本大量重叠在一起，类别边界模糊不清。特别是数字1、7、9等相似数字几乎混为一团，难以区分。这是因为PCA是线性方法，只能捕捉数据的全局方差最大方向，而手写数字的差异更多体现在非线性的局部结构上。')
    
    # t-SNE可视化图
    if os.path.exists(os.path.join(FIG_DIR, 'tsne_2d_visualization.png')):
        add_figure(doc, os.path.join(FIG_DIR, 'tsne_2d_visualization.png'),
                  '图4-3 t-SNE二维降维可视化结果（训练集）')
    
    add_paragraph(doc, '相比之下，图4-3展示的t-SNE可视化效果明显更优。可以清晰地看到10个数字大致形成了10个独立的聚类簇，同类样本聚集紧密，不同类样本之间边界相对清晰。例如数字0、数字6、数字2等都形成了比较集中的团簇。这充分体现了t-SNE作为非线性降维方法在保持局部邻域结构方面的优势，能够更好地揭示数据的内在聚类结构。')
    add_paragraph(doc, '不过t-SNE可视化也存在一些不足：部分相似数字（如数字3和数字8、数字1和数字9）之间仍存在一定程度的重叠；此外，簇与簇之间的距离不具有全局意义，不能直接解读为类别间的真实距离。')
    
    add_heading(doc, '4.3 分类准确率对比分析', level=2)
    add_paragraph(doc, '为了定量评估降维方法对分类性能的影响，本文在原始64维数据、PCA降维2维数据、t-SNE降维2维数据上分别使用三种分类器进行实验。表4-1展示了详细的分类准确率对比结果。')
    
    # 分类准确率表格
    acc_data = [
        ['SVM (RBF核)', '0.9815', '0.5500', '0.9537'],
        ['KNN (k=5)', '0.9704', '0.5148', '0.9704'],
        ['随机森林', '0.9685', '0.5037', '0.9741'],
    ]
    add_table(doc, acc_data, 
             ['分类器', '原始数据(64维)', 'PCA降维(2维)', 't-SNE降维(2维)'],
             '表4-1 不同降维方法下的分类测试准确率对比')
    
    # 准确率对比图
    if os.path.exists(os.path.join(FIG_DIR, 'accuracy_comparison.png')):
        add_figure(doc, os.path.join(FIG_DIR, 'accuracy_comparison.png'),
                  '图4-4 不同降维方法下的分类准确率对比柱状图')
    
    add_paragraph(doc, '从表4-1和图4-4可以得出以下结论：')
    add_paragraph(doc, '第一，原始高维数据上三种分类器均取得了很高的准确率，SVM达到98.15%，说明64维特征包含了足够的区分信息。')
    add_paragraph(doc, '第二，PCA降维到2维后，所有分类器的准确率都出现了大幅下降，均在50%-55%左右。这与方差分析的结果一致——前两个主成分仅保留了约22%的方差信息，大量判别信息在降维过程中丢失。这也说明PCA作为线性全局方法，在降至极低维度时会损失大量对分类有用的信息。')
    add_paragraph(doc, '第三，t-SNE降维到2维后，分类准确率依然保持在很高的水平：随机森林达到97.41%，KNN达到97.04%，SVM达到95.37%。其中KNN和随机森林的准确率甚至与原始64维数据相当或接近，这充分说明t-SNE能够在降维的同时有效保留对分类至关重要的局部邻域结构信息。')
    add_paragraph(doc, '第四，值得注意的是，t-SNE降维后KNN的表现尤为突出，准确率与原始数据完全一致（97.04%）。这是因为KNN本质上依赖局部邻域关系，而t-SNE的优化目标正是保持局部邻域结构，两者高度契合。')
    
    add_heading(doc, '4.4 计算效率对比分析', level=2)
    add_paragraph(doc, '除了分类性能，计算效率也是实际应用中需要考虑的重要因素。表4-2对比了两种降维算法的运行时间。')
    
    time_data = [
        ['PCA (2维)', '0.0042'],
        ['t-SNE (2维)', '2.9067'],
    ]
    add_table(doc, time_data, ['降维方法', '降维耗时（秒）'], '表4-2 降维算法计算时间对比')
    
    # 时间对比图
    if os.path.exists(os.path.join(FIG_DIR, 'time_comparison.png')):
        add_figure(doc, os.path.join(FIG_DIR, 'time_comparison.png'),
                  '图4-5 不同方法的总计算时间对比（SVM分类器）')
    
    add_paragraph(doc, '从计算时间来看，PCA的效率远高于t-SNE。PCA仅需约0.004秒即可完成降维，而t-SNE需要约2.9秒，耗时是PCA的近700倍。这主要有两方面原因：一是PCA基于矩阵分解，有成熟的快速算法（SVD），计算复杂度为O(nd²)；二是t-SNE需要迭代优化，每轮迭代都需要计算所有点对的距离和梯度，计算复杂度为O(n²)，在大数据集上尤为明显。')
    add_paragraph(doc, '此外，t-SNE还有一个重要局限：它没有显式的投影函数，对于新的测试样本无法直接进行降维映射，需要重新训练或使用其他近似方法。而PCA则可以直接将新样本投影到已学习的主成分方向上，更适合在线学习和大规模应用场景。')
    
    doc.add_page_break()
    
    # ========== 5 总结 ==========
    add_heading(doc, '5 总结', level=1)
    add_paragraph(doc, '本文以手写数字识别任务为例，系统对比了线性降维方法PCA和非线性降维方法t-SNE在可视化效果、分类性能和计算效率三个方面的表现。通过实验研究，得出以下主要结论：')
    
    add_paragraph(doc, '第一，在可视化效果方面，t-SNE显著优于PCA。t-SNE能够清晰地展现数据的聚类结构，同类样本聚集紧密，不同类别分离度高；而PCA由于是线性全局方法，降维到二维后各类样本严重重叠，可视化效果较差。', indent=False)
    add_paragraph(doc, '第二，在分类性能方面，当降至相同的2维时，t-SNE远优于PCA。t-SNE降维后分类准确率仍能保持在95%以上，与原始高维数据相差不大；而PCA降维后准确率仅为50%左右，大量判别信息丢失。这说明t-SNE能更好地保留与分类相关的局部结构信息。', indent=False)
    add_paragraph(doc, '第三，在计算效率方面，PCA远优于t-SNE。PCA计算速度快，有显式投影函数，可直接处理新样本；t-SNE计算复杂度高，迭代耗时长，且不支持直接的样本外扩展。', indent=False)
    
    add_paragraph(doc, '综合来看，两种方法各有优劣，适用于不同的应用场景：当数据具有线性结构、需要快速降维、或需要处理新样本时，PCA是更好的选择；当主要目的是数据可视化、探索数据内在结构、且数据量适中时，t-SNE能够提供更优的可视化效果。')
    
    add_paragraph(doc, '本研究也存在一定的局限性：首先，实验仅使用了一个中等规模的数据集，未来可以在更大规模、更高维度的数据集上进行验证；其次，仅对比了两种降维方法，后续可以纳入更多方法如LLE、Isomap、UMAP等进行更全面的比较；最后，t-SNE的超参数（如困惑度）对结果影响较大，未来可以进一步研究超参数的敏感性和调优策略。')
    
    doc.add_page_break()
    
    # ========== 参考文献 ==========
    add_heading(doc, '参考文献', level=1)
    
    references = [
        '[1] Pearson K. On lines and planes of closest fit to systems of points in space[J]. Philosophical Magazine, 1901, 2(11): 559-572.',
        '[2] Hotelling H. Analysis of a complex of statistical variables into principal components[J]. Journal of Educational Psychology, 1933, 24(6): 417-441.',
        '[3] Tenenbaum J B, De Silva V, Langford J C. A global geometric framework for nonlinear dimensionality reduction[J]. Science, 2000, 290(5500): 2319-2323.',
        '[4] Roweis S T, Saul L K. Nonlinear dimensionality reduction by locally linear embedding[J]. Science, 2000, 290(5500): 2323-2326.',
        '[5] Hinton G E, Roweis S T. Stochastic neighbor embedding[C]//Advances in Neural Information Processing Systems. 2002: 833-840.',
        '[6] van der Maaten L, Hinton G. Visualizing data using t-SNE[J]. Journal of Machine Learning Research, 2008, 9(11): 2579-2605.',
        '[7] McInnes L, Healy J, Melville J. UMAP: Uniform manifold approximation and projection for dimension reduction[J]. arXiv preprint arXiv:1802.03426, 2018.',
        '[8] Belkin M, Niyogi P. Laplacian eigenmaps for dimensionality reduction and data representation[J]. Neural Computation, 2003, 15(6): 1373-1396.',
        '[9] Donoho D L, Grimes C. Hessian eigenmaps: Locally linear embedding techniques for high-dimensional data[J]. Proceedings of the National Academy of Sciences, 2003, 100(10): 5591-5596.',
        '[10] Zhang Z, Zha H. Principal manifolds and nonlinear dimensionality reduction via tangent space alignment[J]. SIAM Journal on Scientific Computing, 2004, 26(1): 313-338.',
        '[11] Cox T F, Cox M A A. Multidimensional Scaling[M]. Second edition. London: Chapman & Hall/CRC, 2001.',
        '[12] Fisher R A. The use of multiple measurements in taxonomic problems[J]. Annals of Eugenics, 1936, 7(2): 179-188.',
        '[13] Alpaydin E, Kaynak C. Pen-based recognition of handwritten digits data set[EB/OL]. UCI Machine Learning Repository, 1998.',
        '[14] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.',
        '[15] Cunningham J P, Ghahramani Z. Linear dimensionality reduction: Survey, insights, and generalizations[J]. Journal of Machine Learning Research, 2015, 16(1): 2859-2900.',
        '[16] van der Maaten L. Accelerating t-SNE using tree-based algorithms[J]. Journal of Machine Learning Research, 2014, 15(1): 3221-3245.',
        '[17] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.',
        '[18] 李航. 统计学习方法[M]. 第二版. 北京: 清华大学出版社, 2019.',
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=False, alignment='justify')
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)  # 悬挂缩进
        run = p.add_run(ref)
        set_font(run, font_name='宋体', size=12, bold=False)
    
    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f'论文已生成: {OUTPUT_PATH}')


if __name__ == '__main__':
    generate_paper()
